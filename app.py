import streamlit as st
import json
from PyPDF2 import PdfReader
import docx
from pymongo import MongoClient
import datetime
import hashlib
import time

# ----------------------------------------
# 🔑 ŞİFRELER VE BAĞLANTILAR (Secrets)
# ----------------------------------------
MONGO_URI = st.secrets.get("MONGO_URI", "")

# MongoDB Bağlantısı
@st.cache_resource
def init_mongodb():
    if MONGO_URI:
        try:
            client = MongoClient(MONGO_URI)
            return client["cv_db"]
        except Exception as e:
            st.error(f"Veritabanı bağlantı hatası: {str(e)}")
            return None
    return None

db = init_mongodb()
db_cv_collection = db["cv_collection"] if db is not None else None
db_user_collection = db["users"] if db is not None else None

def hash_password(password):
    return hashlib.sha256(str.encode(password)).hexdigest()

# 1. Dosya Okuma Fonksiyonları
def read_pdf(file):
    pdf_reader = PdfReader(file)
    text = ""
    for page in pdf_reader.pages:
        text += page.extract_text() or ""
    return text

def read_docx(file):
    try:
        doc = docx.Document(file)
        text_parts = []
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        return "\n".join(text_parts)
    except Exception as e:
        return f"Word dosyası okunurken hata oluştu: {str(e)}"

# 2. 🧪 GÜVENLİ TEST MODU FONKSİYONU (Google API Kullanmaz, Kota Harcamaz)
def analyze_cv_with_test_mode(cv_text, filename):
    # Yapay zekayı taklit etmek için 1 saniye bekletme efekti veriyoruz
    time.sleep(1)
    
    # Test amaçlı her yüklenen CV için otomatik üretilecek sahte veri yapısı
    test_sonucu = {
        "Ad Soyad": f"Test Adayı ({filename})",
        "Telefon": "+90 555 123 45 67",
        "E-posta": "aday_test@example.com",
        "Adres": "Kadıköy, İstanbul",
        "Egitim": [
            "Boğaziçi Üniversitesi - Bilgisayar Mühendisliği (2022)",
            "İstanbul Anadolu Lisesi (2017)"
        ],
        "Deneyim": [
            "CSS Ship Management - Teknik Departman Stajyeri (2025 - 2026)",
            "Ceyiznet E-Ticaret - Full Stack Geliştirici (2024 - 2025)"
        ],
        "Yetenekler": "Python, Streamlit, MongoDB, Git, HTML, CSS",
        "Vize Pasaport": "Bordo Pasaport, US Visa (Aktif)",
        "Referanslar": "Süleyman Bey - Genel Müdür (Referans mektubu mevcuttur)"
    }
    return test_sonucu

# 3. Arayüz Tasarımı
st.set_page_config(page_title="AI CV Parser Pro", page_icon="🤖", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #f4f6f9; }
    .stButton>button { width: 100%; border-radius: 8px; font-weight: bold; }
    .result-card { background-color: #ffffff; padding: 20px; border-radius: 10px; box-shadow: 0 4px 6px rgba(0,0,0,0.05); margin-bottom: 15px; }
    .bullet-item { background-color: #f8fafc; padding: 10px 15px; border-left: 4px solid #3b82f6; border-radius: 4px; margin-bottom: 8px; font-size: 14px; }
    </style>
""", unsafe_allow_html=True)

st.title("🤖 Yapay Zeka Destekli CV Analiz Merkezi [TEST MODU]")
st.warning("⚠️ Sistem şu an Ücretsiz Test Modundadır. Google API kotası harcanmaz, sınırsız test yapabilirsiniz.")

if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
if 'user_email' not in st.session_state:
    st.session_state.user_email = ""

# ----------------------------------------
# 🚪 YAN PANEL: ÜYELİK VE PAKET SİSTEMİ
# ----------------------------------------
with st.sidebar:
    if not st.session_state.logged_in:
        st.subheader("🔐 Kullanıcı Paneli")
        auth_mode = st.radio("İşlem Seçin", ["Giriş Yap", "Kayıt Ol"])
        
        email = st.text_input("E-posta Adresi").strip()
        password = st.text_input("Şifre", type="password")
        
        if auth_mode == "Kayıt Ol":
            paket_secimi = st.selectbox("Satın Alınacak Paket", [
                "Başlangıç Paketi (10$ - 100 CV)", 
                "Profesyonel Paket (15$ - 500 CV)", 
                "Kurumsal Paket (25$ - Sınırsız CV)"
            ])
            
            if st.button("Hesap Oluştur (Test Modu - Ücretsiz)"):
                if email and password:
                    if db_user_collection.find_one({"email": email}):
                        st.error("Bu e-posta adresi zaten kayıtlı!")
                    else:
                        if "10$" in paket_secimi:
                            hak = 100
                            p_isim = "Başlangıç"
                        elif "15$" in paket_secimi:
                            hak = 500
                            p_isim = "Profesyonel"
                        else:
                            hak = 10000 # Sınırsız paket koruma sınırı
                            p_isim = "Sınırsız (Kurumsal)"
                            
                        bitis_tarihi = (datetime.datetime.now() + datetime.timedelta(days=30)).strftime("%Y-%m-%d")
                        
                        user_data = {
                            "email": email,
                            "password": hash_password(password),
                            "paket_turu": p_isim,
                            "abonelik_durumu": "aktif",
                            "abonelik_bitis": bitis_tarihi,
                            "kalan_hak": hak
                        }
                        db_user_collection.insert_one(user_data)
                        st.success(f"🎉 {p_isim} aboneliği simüle edildi! Giriş yapabilirsiniz.")
                else:
                    st.warning("Lütfen tüm alanları doldurun.")
                    
        elif auth_mode == "Giriş Yap":
            if st.button("Giriş"):
                user = db_user_collection.find_one({"email": email, "password": hash_password(password)})
                if user:
                    st.session_state.logged_in = True
                    st.session_state.user_email = email
                    st.rerun()
                else:
                    st.error("Hatalı e-posta veya şifre!")
    else:
        st.subheader("👤 Hesap Bilgileri")
        st.write(f"**Kullanıcı:** {st.session_state.user_email}")
        
        user_info = db_user_collection.find_one({"email": st.session_state.user_email})
        if user_info:
            st.write(f"**Paket:** {user_info.get('paket_turu', '-')}")
            st.write(f"**Durum:** {user_info.get('abonelik_durumu', 'Pasif').upper()}")
            st.write(f"**Bitiş Tarihi:** {user_info.get('abonelik_bitis', '-')}")
            
            if user_info.get('paket_turu') == "Sınırsız (Kurumsal)":
                st.write("**Kalan Analiz Hakkı:** Sınırsız ♾️")
            else:
                st.write(f"**Kalan Analiz Hakkı:** {user_info.get('kalan_hak', 0)}")
            
        if st.button("Çıkış Yap"):
            st.session_state.logged_in = False
            st.session_state.user_email = ""
            if 'ai_results' in st.session_state:
                del st.session_state.ai_results
            st.rerun()

# ----------------------------------------
# 🖥️ ANA İÇERİK ALANI
# ----------------------------------------
if not st.session_state.logged_in:
    st.info("ℹ️ Test modunda üyelik ve limit düşme yapılarını denemek için sol panelden kayıt olun ve giriş yapın.")
else:
    col_left, col_right = st.columns([1, 1.5])

    with col_left:
        st.subheader("📁 Özgeçmiş Yükleme")
        uploaded_file = st.file_uploader("Herhangi bir PDF veya Word dosyası seçin (Test Amaçlıdır)", type=["pdf", "docx"])
        
        if uploaded_file is not None:
            raw_text = "Test modunda dosya içeriği okunması simüle ediliyor."
            st.success(f"🔄 {uploaded_file.name} başarıyla belleğe alındı.")
            st.write("---")
            
            if st.button("🚀 Analiz Et (Test Modu)", type="primary"):
                user_info = db_user_collection.find_one({"email": st.session_state.user_email})
                bugun = datetime.datetime.now().strftime("%Y-%m-%d")
                
                if user_info and user_info.get("abonelik_durumu") == "aktif" and user_info.get("abonelik_bitis") >= bugun:
                    if user_info.get("kalan_hak", 0) > 0:
                        with st.spinner("Test Modu: Yapay zeka simüle ediliyor..."):
                            # GOOGLE API YERİNE TEST FONKSİYONUNU ÇAĞIRIYORUZ
                            ai_result = analyze_cv_with_test_mode(raw_text, uploaded_file.name)
                            st.session_state.ai_results = ai_result
                            
                            if db_cv_collection is not None:
                                # CV'yi veritabanına kaydet
                                ai_result["owner_email"] = st.session_state.user_email
                                ai_result["kayit_tarihi"] = datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S")
                                db_cv_collection.insert_one(ai_result)
                                
                                # MongoDB'den hakkı 1 düşür
                                db_user_collection.update_one(
                                    {"email": st.session_state.user_email},
                                    {"$inc": {"kalan_hak": -1}}
                                )
                                st.success("💾 Test verisi başarıyla MongoDB veritabanına kaydedildi ve hakkınız düşüldü!")
                                st.rerun()
                    else:
                        st.error("❌ Bu ayki CV analiz limitinizi (kotanızı) doldurdunuz!")
                else:
                    st.error("❌ Abonelik süreniz dolmuş.")

    with col_right:
        st.subheader("📊 Anlık Sonuç")
        if 'ai_results' in st.session_state:
            res = st.session_state.ai_results
            st.balloons()
            st.markdown('<div class="result-card">', unsafe_allow_html=True)
            st.markdown(f"### 👤 Kişisel Bilgiler")
            st.write(f"**Adı Soyadı:** {res.get('Ad Soyad', 'Belirtilmemiş')}")
            st.write(f"**📞 Telefon:** {res.get('Telefon', 'Belirtilmemiş')}")
            st.write(f"**📧 E-posta:** {res.get('E-posta', 'Belirtilmemiş')}")
            st.write(f"**📍 Adres:** {res.get('Adres', 'Belirtilmemiş')}")
            st.markdown('</div>', unsafe_allow_html=True)
            
            st.markdown('<div class="result-card">', unsafe_allow_html=True)
            st.markdown(f"### 💼 Örnek İş Tecrübeleri")
            for is_madde in res.get('Deneyim', []):
                st.markdown(f'<div class="bullet-item">💼 {is_madde}</div>', unsafe_allow_html=True)
            st.markdown('</div>', unsafe_allow_html=True)
        else:
            st.info("Yüklediğiniz CV'nin simüle edilmiş sonucu burada görünecektir.")

    # KULLANICIYA ÖZEL GEÇMİŞ
    st.write("---")
    st.subheader("🗄️ Geçmiş CV Analizleriniz")

    if db_cv_collection is not None:
        try:
            kayitlar = list(db_cv_collection.find({"owner_email": st.session_state.user_email}).sort("_id", -1))
            
            if len(kayitlar) > 0:
                for kayit in kayitlar:
                    with st.expander(f"📄 {kayit.get('Ad Soyad', 'İsimsiz Aday')} - {kayit.get('kayit_tarihi', '')}"):
                        st.write(f"**📞 Telefon:** {kayit.get('Telefon', '-')}")
                        st.write(f"**📧 E-posta:** {kayit.get('E-posta', '-')}")
                        st.write(f"**💡 Yetenekler:** {kayit.get('Yetenekler', '-')}")
                        st.write("**💼 Deneyimler:**")
                        for d in kayit.get('Deneyim', []):
                            st.write(f"- {d}")
            else:
                st.info("Hesabınıza ait geçmiş bir analiz bulunmuyor.")
        except Exception as e:
            st.error(f"Veriler listelenirken hata oluştu: {str(e)}")
